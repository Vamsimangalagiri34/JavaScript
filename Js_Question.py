from docx import Document

# Create a new Document
doc = Document()

# Add the title
doc.add_heading('JavaScript Questions and Answers', 0)

# List of questions and answers
qa_pairs = [
    ("What are the data types in JavaScript?", "The primitive data types in JavaScript are: \n1. Undefined \n2. Null \n3. Boolean \n4. Number \n5. String \n6. Symbol (ES6) \n7. BigInt (ES11). \nNon-primitive data types are objects, arrays, and functions."),
    ("What is the difference between == and ===?", "The `==` operator compares two values for equality, but it performs type coercion. \nThe `===` operator compares both value and type without type coercion."),
    ("What is the difference between null and undefined?", "`null` is an assignment value, indicating no value. \n`undefined` indicates that a variable has been declared but not yet assigned a value."),
    ("Explain the concept of hoisting in JavaScript.", "Hoisting is JavaScript's default behavior of moving declarations to the top of their scope before code execution. Only the declarations are hoisted, not the assignments."),
    ("What is the difference between let, const, and var?", "`let` and `const` are block-scoped variables, while `var` is function-scoped. \n`const` is for variables that can't be reassigned, while `let` can be reassigned."),
    ("What is variable scope in JavaScript?", "Scope determines the accessibility of variables. JavaScript has three types of scope: global, function, and block scope."),
    ("Explain the difference between global and local variables.", "Global variables are accessible from anywhere in the code, while local variables are accessible only within the function or block where they are declared."),
    ("What is the temporal dead zone?", "The temporal dead zone is the time between the entering of the scope and the actual declaration of the variable, where the variable is in a 'dead' state and cannot be accessed."),
    ("What is variable shadowing?", "Variable shadowing occurs when a local variable or parameter has the same name as a variable in the outer scope, causing the inner variable to overshadow the outer one."),
    ("What is a closure in JavaScript?", "A closure is a function that retains access to its lexical scope, even when the function is executed outside that scope."),
    ("What are the different ways to define a function in JavaScript?", "Functions in JavaScript can be defined in multiple ways: \n1. Function declarations \n2. Function expressions \n3. Arrow functions \n4. Function constructors."),
    ("What is a higher-order function?", "A higher-order function is a function that either takes one or more functions as arguments, returns a function, or both."),
    ("Explain the concept of function hoisting.", "Function hoisting refers to JavaScript’s behavior of moving function declarations to the top of their scope. This allows you to call functions before they appear in the code."),
    ("What is a pure function?", "A pure function is one that always produces the same output for the same input and has no side effects, i.e., it doesn't modify any external state."),
    ("What is the difference between function declaration and function expression?", "Function declarations are hoisted, while function expressions are not. A function expression defines a function inside an expression, while a function declaration defines a function independently."),
    ("What is an Immediately Invoked Function Expression (IIFE)?", "An IIFE is a function that is defined and immediately called (executed) after its creation, typically used to create a new scope in JavaScript."),
    ("How do you create an object in JavaScript?", "Objects in JavaScript can be created using either the object literal syntax `{}` or by using the `new Object()` syntax."),
    ("How do you add/remove properties to an object dynamically?", "Properties can be added or removed from an object using the dot notation (`object.property = value`) or bracket notation (`object['property'] = value`) to add. To remove, you can use the `delete` operator (`delete object.property`)."),
    ("How do you check if a property exists in an object?", "You can check if a property exists in an object using the `in` operator (`'property' in object`), or by using `object.hasOwnProperty('property')`. "),
    ("What is the purpose of the this keyword in JavaScript?", "The `this` keyword refers to the object that is currently executing the code. Its value depends on the context in which the function is called."),
    ("What are the different ways to loop through an array in JavaScript?", "Arrays can be looped through using traditional `for` loops, `for...of`, `forEach()`, `map()`, and other array methods such as `filter()` and `reduce()`."),
    ("Explain the difference between for...in and for...of loops.", "`for...in` iterates over the keys (indices) of an object or array, while `for...of` iterates over the values of an array or iterable object."),
    ("How do you add/remove elements from an array?", "You can use the `push()` and `unshift()` methods to add elements, and `pop()` and `shift()` methods to remove elements from the array."),
    ("What is the purpose of the map() function?", "`map()` is used to apply a function to each item in an array, returning a new array with the results."),
    ("Explain the difference between filter() and find() methods.", "`filter()` returns an array with all elements that pass the test, while `find()` returns the first element that satisfies the condition."),
    ("Explain the difference between some() and every() method.", "`some()` checks if at least one element satisfies the condition, while `every()` checks if all elements satisfy the condition."),
    ("How do you select elements in the DOM using JavaScript?", "DOM elements can be selected using `document.getElementById()`, `document.getElementsByClassName()`, `document.querySelector()`, and other selection methods."),
    ("How do you create and append elements to the DOM?", "To create and append an element, use `document.createElement()` to create a new element and `parentElement.appendChild()` to add it to the DOM."),
    ("Explain the difference between innerHTML and textContent.", "`innerHTML` gets or sets the HTML content of an element, while `textContent` gets or sets only the text content, excluding any HTML tags."),
    ("How do you remove an element from the DOM?", "To remove an element, use `parentElement.removeChild(childElement)` or the `remove()` method directly on the element."),
    ("What are arrow functions and how do they differ from regular functions?", "Arrow functions are shorthand function expressions that use `=>`. They do not have their own `this`, `arguments`, or `super` and are best used in simple scenarios."),
    ("Explain the concept of destructuring in JavaScript.", "Destructuring allows you to unpack values from arrays or properties from objects into distinct variables."),
    ("What are template literals?", "Template literals are string literals allowing embedded expressions. They are enclosed by backticks (` `) and can contain placeholders using `${expression}`)."),
    ("How do you use the spread operator?", "The spread operator (`...`) is used to expand elements of an iterable (array, object) into individual elements."),
    ("What are default parameters in ES6?", "Default parameters in ES6 allow you to set default values for function parameters, which will be used if no argument is passed."),
    ("How do you use the rest parameter in functions?", "The rest parameter (`...`) allows a function to accept an indefinite number of arguments as an array."),
    ("What is callback & callback hell explain with example", "A callback is a function passed into another function as an argument. Callback hell happens when callbacks are nested in multiple layers, making the code hard to read and manage."),
    ("What is a Promise in JavaScript with example?", "A Promise represents an asynchronous operation. It can be in one of three states: pending, resolved, or rejected. Example: `let promise = new Promise((resolve, reject) => { /* logic */ });`"),
    ("How do you chain Promises?", "You can chain promises using the `.then()` method, where each `.then()` returns a new promise, allowing for sequential asynchronous operations."),
    ("What is the purpose of the Promise.all() method?", "`Promise.all()` is used to execute multiple promises in parallel and wait for all of them to resolve or for any to reject."),
    ("What is the purpose of the finally() method in Promises?", "`finally()` executes code after a promise is resolved or rejected, regardless of the outcome."),
    ("What is the purpose of the async await?", "`async` and `await` are used to work with promises in a more synchronous manner, making asynchronous code easier to read and write."),
    ("How do you handle errors in async/await?", "Errors in `async/await` can be caught using `try...catch` blocks."),
    ("What is the difference between async/await and Promises?", "`async/await` is syntactic sugar built on top of promises, making asynchronous code look and behave more like synchronous code."),
    ("What is the difference between default and named exports?", "Default exports allow a single value to be exported from a module, while named exports allow multiple values to be exported from the same module.")
]

# Adding questions and answers to the document
for question, answer in qa_pairs:
    doc.add_heading(question, level=1)
    doc.add_paragraph(answer)

# Save the document
file_path =Js_question
